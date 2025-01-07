from flask import Flask, request, render_template, redirect, url_for, send_from_directory
import os
import pandas as pd
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import zipfile
from docx import Document  # Import for saving text data to a DOCX file
import tabula  # For PDF to CSV conversion

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

@app.route('/', methods=['GET'])
def welcome():
    return render_template('welcome.html')

@app.route('/upload', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            return redirect(request.url)
        file = request.files['pdf_file']
        if file.filename == '':
            return redirect(request.url)
        if file:
            filepath = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(filepath)
            try:
                text_data, image_data, marks_data = extract_text_and_images_from_pdf(filepath)
                save_to_files(text_data, image_data, marks_data)
                convert_pdf_to_csv(filepath)
            except Exception as e:
                return str(e)
            return redirect(url_for('download_files'))
    return render_template('index.html')

def extract_text_and_images_from_pdf(filepath):
    text_data = []
    image_data = []
    marks_data = []
    try:
        doc = fitz.open(filepath)
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text = page.get_text()
            text_data.append({"Page": page_num + 1, "Text": text})

            image_list = page.get_images(full=True)
            for img_num, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_data.append({"Page": page_num + 1, "Image Number": img_num + 1, "Image": image_bytes})

                image = Image.open(io.BytesIO(image_bytes))
                marks = detect_marks(image)
                marks_data.append({"Page": page_num + 1, "Image Number": img_num + 1, "Marks": marks})
        doc.close()
    except Exception as e:
        raise RuntimeError(f"Error extracting text and images from PDF: {e}")
    return text_data, image_data, marks_data

def detect_marks(image):
    gray_image = image.convert('L')
    ocr_result = pytesseract.image_to_string(gray_image)
    
    marks = {
        'tick': '✓' in ocr_result,
        'cross': '✗' in ocr_result,
        'shaded_box': '■' in ocr_result or '▇' in ocr_result
    }
    
    return marks

def save_to_files(text_data, image_data, marks_data):
    try:
        # Save text data to a text file
        text_file_path = os.path.join(OUTPUT_FOLDER, 'extracted_text.txt')
        with open(text_file_path, 'w', encoding='utf-8') as text_file:
            for item in text_data:
                text_file.write(f"Page {item['Page']}:\n{item['Text']}\n\n")

        # Save text data to a DOCX file
        doc = Document()
        for item in text_data:
            doc.add_heading(f"Page {item['Page']}", level=1)
            doc.add_paragraph(item['Text'])
        doc_path = os.path.join(OUTPUT_FOLDER, 'extracted_text.docx')
        doc.save(doc_path)

        # Convert text data to DataFrame and save as Excel and CSV
        text_df = pd.DataFrame(text_data)
        excel_file_path = os.path.join(OUTPUT_FOLDER, 'extracted_text.xlsx')
        csv_file_path = os.path.join(OUTPUT_FOLDER, 'extracted_text.csv')
        text_df.to_excel(excel_file_path, index=False)
        text_df.to_csv(csv_file_path, index=False)

        # Save images
        image_folder_path = os.path.join(OUTPUT_FOLDER, 'images')
        if not os.path.exists(image_folder_path):
            os.makedirs(image_folder_path)
        for image in image_data:
            image_file_path = os.path.join(image_folder_path, f"page_{image['Page']}_image_{image['Image Number']}.png")
            with open(image_file_path, 'wb') as image_file:
                image_file.write(image["Image"])

        # Save detected marks
        marks_file_path = os.path.join(OUTPUT_FOLDER, 'detected_marks.txt')
        with open(marks_file_path, 'w', encoding='utf-8') as marks_file:
            for item in marks_data:
                marks_file.write(f"Page {item['Page']}, Image {item['Image Number']}:\n{item['Marks']}\n\n")

        marks_df = pd.DataFrame(marks_data)
        marks_excel_file_path = os.path.join(OUTPUT_FOLDER, 'detected_marks.xlsx')
        marks_csv_file_path = os.path.join(OUTPUT_FOLDER, 'detected_marks.csv')
        marks_df.to_excel(marks_excel_file_path, index=False)
        marks_df.to_csv(marks_csv_file_path, index=False)
    except Exception as e:
        raise RuntimeError(f"Error saving files: {e}")

def convert_pdf_to_csv(filepath):
    # Convert PDF to CSV
    output_csv_path = os.path.join(OUTPUT_FOLDER, 'converted.csv')
    try:
        tabula.convert_into(filepath, output_csv_path, output_format='csv', pages='all')
    except Exception as e:
        raise RuntimeError(f"Error converting PDF to CSV: {e}")

@app.route('/download')
def download_files():
    return render_template('download.html')

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(OUTPUT_FOLDER, filename)

@app.route('/download/images')
def download_images():
    # Path to the images folder
    image_folder_path = os.path.join(OUTPUT_FOLDER, 'images')

    # Path to the zip file
    zip_filename = 'images.zip'
    zip_file_path = os.path.join(OUTPUT_FOLDER, zip_filename)

    # Create a zip file of the images folder
    with zipfile.ZipFile(zip_file_path, 'w') as zipf:
        for root, dirs, files in os.walk(image_folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                # Add file to the zip file, keeping folder structure
                zipf.write(file_path, os.path.relpath(file_path, image_folder_path))

    return send_from_directory(OUTPUT_FOLDER, zip_filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)

